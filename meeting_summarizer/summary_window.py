from PyQt6.QtWidgets import (QDialog, QWidget, QVBoxLayout, QLabel, QTextEdit, QPushButton, 
                            QComboBox, QFrame, QHBoxLayout, QFileDialog, QApplication, 
                            QMessageBox, QGroupBox)
from PyQt6.QtCore import Qt
from PyQt6.QtGui import QFont
from utils.MeetingRecordProject import MeetingRecordProject
import os
from docx import Document
from utils.flexible_logger import Logger
import traceback

class SummaryWindow(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.logger = Logger(
            name="summary_window",
            console_output=True,
            file_output=True,
            log_level="INFO"
        )
        self.project_manager = None
        
        # 设置窗口属性
        self.setWindowTitle("会议总结")
        self.setModal(True)  # 设置为模态对话框
        self.resize(800, 600)  # 设置合适的窗口大小
        
        self.init_ui()
        self.logger.info("SummaryWindow initialized")
    
    def init_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(20, 20, 20, 20)
        main_layout.setSpacing(20)
        
        # Summary text area
        self.summary_text = QTextEdit()
        self.summary_text.setReadOnly(True)
        self.summary_text.setMinimumHeight(400)  # 增加高度
        self.summary_text.setStyleSheet("""
            QTextEdit {
                border: 1px solid #CCCCCC;
                border-radius: 5px;
                padding: 10px;
                background-color: white;
            }
        """)
        main_layout.addWidget(self.summary_text)
        
        # Export options group
        export_group = QGroupBox()
        export_layout = QHBoxLayout()
        export_layout.setSpacing(20)
        
        # Language selection
        language_label = QLabel("输出语言:")
        self.language_combo = QComboBox()
        self.language_combo.addItems(["中文", "English"])
        self.language_combo.setFixedWidth(150)
        self.language_combo.setStyleSheet("""
            QComboBox {
                border: 2px solid #CCCCCC;
                border-radius: 5px;
                padding: 5px 10px;
                background: white;
                font-size: 12px;
                min-height: 25px;
            }
            QComboBox:hover {
                border: 2px solid #33CCFF;
            }
            QComboBox:focus {
                border: 2px solid #33CCFF;
            }
            QComboBox::drop-down {
                border: none;
                width: 20px;
                padding-right: 5px;
            }
            QComboBox::down-arrow {
                width: 12px;
                height: 12px;
                image: url(resources/icons/down-arrow.png);
            }
            QComboBox QAbstractItemView {
                border: 1px solid #CCCCCC;
                selection-background-color: #33CCFF;
                selection-color: white;
                background: white;
            }
        """)
        export_layout.addWidget(language_label)
        export_layout.addWidget(self.language_combo)
        
        # Format selection
        format_label = QLabel("输出格式:")
        self.format_combo = QComboBox()
        self.format_combo.addItems(["Markdown", "Word"])
        self.format_combo.setFixedWidth(150)
        self.format_combo.setStyleSheet(self.language_combo.styleSheet())  # 使用相同的样式
        export_layout.addWidget(format_label)
        export_layout.addWidget(self.format_combo)
        
        # Export button
        self.export_button = QPushButton("导出")
        self.export_button.clicked.connect(self.export_summary)
        self.export_button.setStyleSheet("""
            QPushButton {
                background-color: #33CCFF;
                color: white;
                border: none;
                border-radius: 5px;
                padding: 15px 30px;
                font-size: 16px;
            }
            QPushButton:hover {
                background-color: #2CB5E8;
            }
            QPushButton:disabled {
                background-color: #CCCCCC;
            }
        """)
        export_layout.addWidget(self.export_button)
        
        export_layout.addStretch()  # 添加弹性空间
        export_group.setLayout(export_layout)
        main_layout.addWidget(export_group)
        
        # Bottom buttons
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        
        # Close button
        close_button = QPushButton("关闭")
        close_button.clicked.connect(self.accept)
        close_button.setStyleSheet("""
            QPushButton {
                background-color: #6C757D;
                color: white;
                border: none;
                border-radius: 5px;
                padding: 15px 30px;
                font-size: 16px;
            }
            QPushButton:hover {
                background-color: #5A6268;
            }
        """)
        button_layout.addWidget(close_button)
        
        main_layout.addLayout(button_layout)
    
    def set_project_manager(self, project_manager):
        """Set project manager"""
        try:
            self.project_manager = project_manager
            self.logger.info(f"已设置项目管理器: {project_manager}")
        except Exception as e:
            self.logger.error(f"设置项目管理器失败: {str(e)}")
    
    def load_summary(self):
        """Load summary content from project"""
        try:
            if not self.project_manager:
                raise ValueError("未设置项目管理器")
            
            summary_file = self.project_manager.get_summary_filename()
            if not summary_file or not os.path.exists(summary_file):
                self.summary_text.setText("未找到总结文件")
                return
            
            with open(summary_file, 'r', encoding='utf-8') as f:
                summary_text = f.read()
                self.summary_text.setPlainText(summary_text)
                
        except Exception as e:
            self.logger.error(f"加载总结文件时发生错误: {str(e)}")
            traceback.print_exc() 
    
    def show_summary(self):
        """显示总结对话框并加载内容"""
        self.load_summary()
        self.exec()  # 使用 exec() 显示模态对话框
    
    def export_summary(self):
        """导出总结内容"""
        try:
            # 获取当前选择的语言和格式
            language = self.language_combo.currentText()
            format_type = self.format_combo.currentText()
            self.logger.info(f"导出总结: 语言: {language}, 格式: {format_type}")

            # 获取要导出的内容
            content = self.summary_text.toPlainText()
            if not content:
                QMessageBox.warning(self, "导出失败", "没有可导出的内容")
                return
            
            # 选择保存路径
            default_name = f"会议总结_{language}"
            file_extension = ".md" if format_type == "Markdown" else ".docx"
            self.logger.info(f"选择保存路径: {default_name + file_extension}")
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "保存总结",
                default_name + file_extension,
                f"{'Markdown 文件 (*.md)' if format_type == 'Markdown' else 'Word 文档 (*.docx)'}"
            )
            self.logger.info(f"选择保存路径: {file_path}")
            
            if not file_path:
                return
            
            # 根据不同格式导出
            if format_type == "Markdown":
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
            else:  # Word format
                doc = Document()
                doc.add_heading('会议总结', 0)
                doc.add_paragraph(content)
                doc.save(file_path)
            
            self.logger.info(f"总结已导出到: {file_path}")
            QMessageBox.information(self, "导出成功", f"总结已导出到:\n{file_path}")
            
        except Exception as e:
            self.logger.error(f"导出总结时发生错误: {str(e)}")
            QMessageBox.critical(self, "导出失败", f"导出过程中发生错误:\n{str(e)}")
            traceback.print_exc()