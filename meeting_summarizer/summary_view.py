from PyQt6.QtWidgets import (QWidget, QVBoxLayout, QLabel, QTextEdit, QPushButton, 
                            QComboBox, QFrame, QHBoxLayout, QFileDialog, QApplication, QMessageBox, QGroupBox)
from PyQt6.QtCore import Qt
from PyQt6.QtGui import QFont
from utils.MeetingRecordProject import MeetingRecordProject
import os
from docx import Document
from utils.llm_proofreader import TextProofreader
from utils.lecture_notes_generator import LectureNotesGenerator
from utils.meeting_notes_generator import MeetingNotesGenerator
from utils.flexible_logger import Logger
import traceback

class SummaryViewWidget(QWidget):
    def __init__(self):
        super().__init__()
        self.logger = Logger(
            name="summary_view",
            console_output=True,
            file_output=True,
            log_level="INFO"
        )
        self.project_manager = None
        self.proofreader = TextProofreader()  # 初始化校对器
        self.current_summary_file = None
        self.init_ui()
        self.logger.info("SummaryViewWidget initialized")
        
    def init_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(20, 20, 20, 20)
        main_layout.setSpacing(20)
        
        # 顶部布局（包含标题、LLM状态和导出选项）
        top_layout = QHBoxLayout()
        
        # 左侧标题
        title = QLabel("会议纪要预览")
        title.setFont(QFont("Arial", 24, QFont.Weight.Bold))
        top_layout.addWidget(title)
        
        # 添加LLM状态按钮
        self.llm_status_button = QPushButton()
        self.llm_status_button.setStyleSheet("""
            QPushButton {
                border: none;
                padding: 5px 10px;
                border-radius: 3px;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #F0F0F0;
            }
        """)
        self.llm_status_button.clicked.connect(self.open_llm_settings)
        top_layout.addWidget(self.llm_status_button)
        
        # 右侧导出选项
        export_group = QFrame()
        export_group.setStyleSheet("""
            QFrame {
                background-color: #F8F8F8;
                border: 1px solid #E0E0E0;
                border-radius: 5px;
                padding: 5px;
            }
        """)
        export_layout = QHBoxLayout(export_group)
        export_layout.setContentsMargins(10, 5, 10, 5)
        export_layout.setSpacing(10)
        
        # 导出语种选择
        lang_label = QLabel("语种：")
        self.lang_combo = QComboBox()
        self.lang_combo.addItems(["中文", "English", "日本語"])
        self.lang_combo.setFixedWidth(100)
        self.lang_combo.setStyleSheet("""
            QComboBox {
                border: 1px solid #E0E0E0;
                border-radius: 3px;
                padding: 3px;
                background: white;
            }
        """)
        
        # 导出格式选择
        format_label = QLabel("格式：")
        self.format_combo = QComboBox()
        self.format_combo.addItems(["Markdown", "PDF", "Word"])
        self.format_combo.setFixedWidth(100)
        self.format_combo.setStyleSheet("""
            QComboBox {
                border: 1px solid #E0E0E0;
                border-radius: 3px;
                padding: 3px;
                background: white;
            }
        """)
        
        # 导出按钮
        export_button = QPushButton("导出")
        export_button.setStyleSheet("""
            QPushButton {
                background-color: #33CCFF;
                color: white;
                border: none;
                border-radius: 3px;
                padding: 5px 15px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2CB5E8;
            }
        """)
        
        # 添加到导出布局
        export_layout.addWidget(lang_label)
        export_layout.addWidget(self.lang_combo)
        export_layout.addWidget(format_label)
        export_layout.addWidget(self.format_combo)
        export_layout.addWidget(export_button)
        
        # 添加模板类型选择
        template_group = QGroupBox("模板类型")
        template_layout = QHBoxLayout(template_group)
        
        template_label = QLabel("选择模板：")
        self.template_combo = QComboBox()
        self.template_combo.addItems(["会议纪要", "课堂笔记"])
        
        template_layout.addWidget(template_label)
        template_layout.addWidget(self.template_combo)
        
        # 添加总结按钮
        self.summarize_button = QPushButton("总结")
        self.summarize_button.clicked.connect(self.generate_summary)
        template_layout.addWidget(self.summarize_button)
        
        # 将模板选择组添加到顶部布局
        top_layout.addWidget(template_group)
        top_layout.addStretch()  # 添加弹性空间
        top_layout.addWidget(export_group)
        
        main_layout.addLayout(top_layout)
        
        # 中间内容区域
        content_layout = QHBoxLayout()
        
        # 左侧会议记录
        transcript_container = QWidget()
        transcript_layout = QVBoxLayout(transcript_container)
        transcript_label = QLabel("会议记录")
        transcript_label.setFont(QFont("Arial", 16))
        self.transcript_text = QTextEdit()
        self.transcript_text.setReadOnly(True)
        transcript_layout.addWidget(transcript_label)
        transcript_layout.addWidget(self.transcript_text)
        
        # 添加校对按钮
        self.proofread_button = QPushButton("校对")
        self.proofread_button.clicked.connect(self.proofread_transcript)
        transcript_layout.addWidget(self.proofread_button)
        
        content_layout.addWidget(transcript_container)
        
        # 添加垂直分割线
        line = QFrame()
        line.setFrameShape(QFrame.Shape.VLine)
        line.setFrameShadow(QFrame.Shadow.Sunken)
        content_layout.addWidget(line)
        
        # 右侧会议纪要
        summary_container = QWidget()
        summary_layout = QVBoxLayout(summary_container)
        
        # 添加标题和语种选择
        header_layout = QHBoxLayout()
        summary_label = QLabel("会议纪要")
        summary_label.setFont(QFont("Arial", 16))
        header_layout.addWidget(summary_label)
        
        # 添加导出语种选择
        lang_layout = QHBoxLayout()
        lang_label = QLabel("导出语种：")
        self.lang_combo = QComboBox()
        self.lang_combo.addItems(["中文", "English", "日本語"])
        lang_layout.addWidget(lang_label)
        lang_layout.addWidget(self.lang_combo)
        header_layout.addLayout(lang_layout)
        header_layout.addStretch()
        
        summary_layout.addLayout(header_layout)
        
        self.summary_text = QTextEdit()
        self.summary_text.setReadOnly(True)
        self.summary_text.setStyleSheet("""
            QTextEdit {
                background-color: white;
                border: 1px solid #E0E0E0;
                border-radius: 5px;
                padding: 10px;
            }
        """)
        summary_layout.addWidget(self.summary_text)
        
        # 操作按钮
        button_layout = QHBoxLayout()
        
        self.edit_button = QPushButton("编辑")
        self.edit_button.clicked.connect(self.toggle_edit)
        button_layout.addWidget(self.edit_button)
        
        self.save_button = QPushButton("保存")
        self.save_button.clicked.connect(self.save_summary)
        self.save_button.setEnabled(False)
        button_layout.addWidget(self.save_button)
        
        self.export_button_summary = QPushButton("导出")
        self.export_button_summary.clicked.connect(self.export_summary)
        button_layout.addWidget(self.export_button_summary)
        
        summary_layout.addLayout(button_layout)
        content_layout.addWidget(summary_container)
        
        main_layout.addLayout(content_layout)
        
        # 连接信号
        export_button.clicked.connect(self.export_summary)
        
        # 初始检查LLM状态
        self.check_llm_status()
    
    def load_content(self):
        """加载所有必要的内容"""
        try:
            self.logger.info("开始加载总结视图内容")
            if not self.project_manager:
                raise ValueError("未设置项目管理器，无法加载内容")
            
            self.logger.info("开始加载转写文件...")
            self.load_transcriptfile()
            self.logger.info("开始加载总结文件...")
            self.load_summaryfile()
            self.logger.info("内容加载完成")
            
        except Exception as e:
            self.logger.error(f"加载内容时发生错误: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def load_latest_summary(self):
        """加载最新的会议总结"""
        try:
            if not self.project_manager:
                self.summary_text.setText("未设置项目管理器")
                return

            summary_file = self.project_manager.get_summary_filename()
            if not summary_file or not os.path.exists(summary_file):
                self.summary_text.setText("未找到会议总结文件")
                return
            
            # 读取总结内容
            with open(summary_file, 'r', encoding='utf-8') as f:
                summary_content = f.read()
            
            self.summary_text.setText(summary_content)
            self.current_summary_file = summary_file
            
        except Exception as e:
            self.logger.error(f"加载总结文件时发生错误: {str(e)}")
            import traceback
            traceback.print_exc()
            self.summary_text.setText(f"加载总结文件失败: {str(e)}")
    
    def load_transcriptfile(self):
        """加载转写文件"""
        try:
            if not self.project_manager:
                raise ValueError("未设置项目管理器")
            
            transcript_file = self.project_manager.get_transcript_filename()
            if not transcript_file or not os.path.exists(transcript_file):
                self.transcript_text.setText("未找到转写文件")
                return
            
            with open(transcript_file, 'r', encoding='utf-8') as f:
                transcript_text = f.read()
                self.transcript_text.setPlainText(transcript_text)
            
        except Exception as e:
            self.logger.error(f"加载转写文件时发生错误: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def load_summaryfile(self):
        """加载总结文件"""
        try:
            if not self.project_manager:
                self.summary_text.setText("未设置项目管理器")
                return

            summary_file = self.project_manager.get_summary_filename()
            if not summary_file or not os.path.exists(summary_file):
                self.summary_text.setText("未找到总结文件")
                return
            
            with open(summary_file, 'r', encoding='utf-8') as f:
                content = f.read()
                self.summary_text.setPlainText(content)
                
        except Exception as e:
            self.logger.error(f"加载总结文件时发生错误: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def toggle_edit(self):
        """切换编辑模式"""
        if self.summary_text.isReadOnly():
            self.summary_text.setReadOnly(False)
            self.edit_button.setText("取消")
            self.save_button.setEnabled(True)
        else:
            self.summary_text.setReadOnly(True)
            self.edit_button.setText("编辑")
            self.save_button.setEnabled(False)
            # 重新加载原内容
            self.load_latest_summary()
    
    def save_summary(self):
        """保存编辑后的总结"""
        try:
            if not self.project_manager:
                raise ValueError("未设置项目管理器")

            # 获取新的总结文件名
            summary_file = self.project_manager.get_summary_new_filename()
            
            # 保存内容
            with open(summary_file, 'w', encoding='utf-8') as f:
                f.write(self.summary_text.toPlainText())
            
            # 更新当前文件
            self.current_summary_file = summary_file
            
            # 更新界面状态
            self.summary_text.setReadOnly(True)
            self.edit_button.setText("编辑")
            self.save_button.setEnabled(False)
            
        except Exception as e:
            self.logger.error(f"保存总结文件时发生错误: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def export_summary(self):
        """导出会议总结"""
        try:
            if hasattr(self, 'current_summary_file'):
                if self.current_summary_file.endswith('.txt'):
                    # 直接复制文本文件
                    import shutil
                    shutil.copy2(self.current_summary_file, self.current_summary_file)
                elif self.current_summary_file.endswith('.docx'):
                    # 转换为Word文档
                    doc = Document()
                    doc.add_heading('会议纪要', 0)
                    doc.add_paragraph(self.summary_text.toPlainText())
                    doc.save(self.current_summary_file)
        except Exception as e:
            self.logger.error(f"导出总结文件时发生错误: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def proofread_transcript(self):
        """处理校对请求"""
        try:
            self.logger.info("开始校对文本...")
            self.proofread_button.setEnabled(False)
            self.proofread_button.setText("校对中...")
            
            # 获取当前文本
            current_text = self.transcript_text.toPlainText()
            if not current_text.strip():
                error_msg = "没有可用的文本内容"
                self.logger.warning(error_msg)
                QMessageBox.warning(self, "错误", error_msg)
                return
            
            try:
                # 发送文本进行校对
                self.logger.info("正在进行文本校对...")
                results = self.proofreader.proofread_text(current_text)
                
                if results and 'proofread_text' in results:
                    # 更新文本框内容
                    self.transcript_text.setPlainText(results['proofread_text'])
                    
                    # 保存校对后的文本
                    if self.project_manager:
                        try:
                            new_transcript_file = self.project_manager.get_transcript_new_filename()
                            self.logger.info(f"保存校对后的文本... {new_transcript_file}")
                            with open(new_transcript_file, 'w', encoding='utf-8') as f:
                                f.write(results['proofread_text'])
                            self.project_manager.add_proofread_transcript(new_transcript_file)
                            self.logger.info(f"校对文本已保存至: {new_transcript_file}")
                        except Exception as save_error:
                            error_msg = f"保存校对文本时发生错误: {str(save_error)}"
                            self.logger.error(error_msg)
                            QMessageBox.warning(self, "保存失败", error_msg)
                else:
                    error_msg = "校对结果格式错误"
                    self.logger.error(error_msg)
                    QMessageBox.warning(self, "校对失败", error_msg)
                    
            except Exception as e:
                error_msg = f"校对过程中发生错误：{str(e)}"
                self.logger.error(error_msg)
                self.logger.debug(f"错误详情: {traceback.format_exc()}")
                QMessageBox.warning(self, "校对失败", error_msg)
                
        except Exception as e:
            error_msg = f"校对功能发生错误: {str(e)}"
            self.logger.error(error_msg)
            self.logger.debug(f"错误详情: {traceback.format_exc()}")
            QMessageBox.warning(self, "错误", error_msg)
            
        finally:
            # 恢复校对按钮状态
            self.proofread_button.setEnabled(True)
            self.proofread_button.setText("校对")
            self.logger.info("校对操作结束")
    
    def set_transcript(self, text):
        """设置会议记录文本"""
        self.transcript_text.setPlainText(text)
    
    def set_summary(self, text):
        """设置会议纪要文本"""
        self.summary_text.setPlainText(text)
    
    def set_project_manager(self, project_manager):
        """设置项目管理器"""
        try:
            self.project_manager = project_manager
            self.logger.info(f"已设置项目管理器: {project_manager}")
        except Exception as e:
            self.logger.error(f"设置项目管理器失败: {str(e)}")
    
    def check_llm_status(self):
        """检查LLM服务状态"""
        try:
            from utils.llm_statuscheck import LLMStatusChecker
            from config.settings import Settings
            
            settings = Settings()
            llm_config = settings._settings["llm"]
            checker = LLMStatusChecker(llm_config["api_url"])
            status = checker.check_status()
            
            if status == "ready":
                self.llm_status_button.setText("LLM服务正常")
                self.llm_status_button.setStyleSheet("""
                    QPushButton {
                        color: white;
                        background-color: #4CAF50;
                        border: none;
                        padding: 5px 10px;
                        border-radius: 3px;
                        font-size: 12px;
                    }
                    QPushButton:hover {
                        background-color: #45a049;
                    }
                """)
            else:
                self.llm_status_button.setText("LLM服务离线")
                self.llm_status_button.setStyleSheet("""
                    QPushButton {
                        color: white;
                        background-color: #f44336;
                        border: none;
                        padding: 5px 10px;
                        border-radius: 3px;
                        font-size: 12px;
                    }
                    QPushButton:hover {
                        background-color: #da190b;
                    }
                """)
        except Exception as e:
            self.logger.error(f"检查LLM状态时发生错误: {str(e)}")
            self.llm_status_button.setText("LLM状态未知")
            self.llm_status_button.setStyleSheet("""
                QPushButton {
                    color: white;
                    background-color: #808080;
                    border: none;
                    padding: 5px 10px;
                    border-radius: 3px;
                    font-size: 12px;
                }
                QPushButton:hover {
                    background-color: #666666;
                }
            """)

    def open_llm_settings(self):
        """打开LLM设置页面"""
        # TODO: 实现打开设置页面的功能
        # 这里需要调用主窗口的方法来打开设置页面
        print("TODO: 打开LLM设置页面")

    def generate_summary(self):
        """生成总结"""
        try:
            # 获取当前文本内容
            text = self.transcript_text.toPlainText()
            if not text.strip():
                QMessageBox.warning(self, "错误", "没有可用的文本内容")
                return
            
            # 根据模板类型选择生成器
            template_type = self.template_combo.currentText()
            
            if template_type == "课堂笔记":
                generator = LectureNotesGenerator()
            else:  # 会议纪要
                generator = MeetingNotesGenerator()
            
            # 生成总结
            try:
                summary = generator.generate_notes(text)
                self.summary_text.setPlainText(summary)
                
                # 如果项目管理器存在，保存总结
                if self.project_manager:
                    summary_file = self.project_manager.get_summary_new_filename()
                    with open(summary_file, 'w', encoding='utf-8') as f:
                        f.write(summary)
                    self.current_summary_file = summary_file
                    
            except Exception as e:
                QMessageBox.warning(
                    self,
                    "总结生成失败",
                    f"生成总结时发生错误：{str(e)}"
                )
                
        except Exception as e:
            self.logger.error(f"生成总结时发生错误: {str(e)}")
            import traceback
            traceback.print_exc()
            QMessageBox.warning(
                self,
                "错误",
                f"生成总结时发生错误：{str(e)}"
            )